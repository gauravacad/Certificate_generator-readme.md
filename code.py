import pandas as pd
from docx import Document
import os
import re

from docx.shared import Pt
from docx.oxml.ns import qn

# === Configuration ===
template_path = "template.docx"
excel_path = "data.xlsx"
output_folder = "output_docs/"

os.makedirs(output_folder, exist_ok=True)

# === Load Excel ===
df = pd.read_excel(excel_path)
df.columns = df.columns.str.strip()  # remove spaces in column headers

'''# === Replace placeholders in a paragraph or cell ===
def replace_placeholders_text(obj, replacements):
    # Join all runs to get full text
    full_text = "".join(run.text for run in obj.runs)

    # Replace all placeholders
    for key, val in replacements.items():
        pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
        full_text = re.sub(pattern, str(val), full_text)
      

    #Clear old runs
    for run in obj.runs:
        run.text = ""

    # Add a run if none exist
    if len(obj.runs) == 0:
        p=obj.add_run(full_text)
        
    else:
        obj.runs[0].text = full_text'''

from docx.shared import Pt

def replace_placeholders_text(paragraph, replacements):
    # Join all runs to get full paragraph text
    full_text = "".join(run.text for run in paragraph.runs)
    
    # If no placeholders, do nothing
    if not any(f"{{{{{k}}}}}" in full_text for k in replacements):
        return

    # Save the formatting from the first run (to preserve template style)
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font_name = first_run.font.name or "Mangal (Body CS)"
        font_size = Pt(12)
    else:
        font_name = "Mangal (Body CS)"
        font_size = Pt(12)

    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    pos = 0
    while pos < len(full_text):
        # Find the next placeholder
        match = None
        for key in replacements:
            pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
            match = re.search(pattern, full_text[pos:])
            if match:
                match_key = key
                break
        if not match:
            # Add remaining text as normal
            r = paragraph.add_run(full_text[pos:])
            r.font.name = font_name
            r.font.size = font_size
            break
        start, end = match.span()
        start += pos
        end += pos

        # Text before placeholder
        if start > pos:
            r = paragraph.add_run(full_text[pos:start])
            r.font.name = font_name
            r.font.size = font_size

        # Replacement text
        replacement_text = str(replacements[match_key])
        r = paragraph.add_run(replacement_text)
        r.font.name = font_name
        r.font.size = font_size
        r.bold = True
        # remove underline in Name
        #if match_key.upper() == "NAME":
        #    r.underline = True
        if match_key.upper() == "POST":
            r.bold = True

        pos = end



# === Replace placeholders in paragraphs and tables ===
def replace_placeholders(doc, replacements):
    for para in doc.paragraphs:
        replace_placeholders_text(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_text(para, replacements)

# === Process each row ===
for _, row in df.iterrows():
    replacements = {col: row[col] for col in df.columns}
    
    doc = Document(template_path)
    replace_placeholders(doc, replacements)
    
    id_ = str(row.get("ID", "0000"))
    name = str(row.get("NAME", "unknown"))
    safe_name = "".join(c for c in name if c.isalnum() or c in (" ", "_")).rstrip()
    output_path = os.path.join(output_folder, f"{id_}_{safe_name}.docx")
    
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")
